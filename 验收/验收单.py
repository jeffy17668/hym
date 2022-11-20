import docx
import pandas as pd
from docx import document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import os
from copy import deepcopy
import time
current_file_path = os.getcwd()
os.chdir(current_file_path)
warnings.filterwarnings('ignore')

start_time = time.time()
filePath1 = 'excel'
file_name1 = os.listdir(filePath1)
for i in range(len(file_name1)):
    if str(file_name1[i]).count('~$') == 0:
        exc = pd.read_excel(filePath1 + '/' + str(file_name1[i]), dtype=object)

# exc=pd.read_excel(r"excel\江海九月要做验收单的.xls", dtype=object)
doc_name = r"验收单.docx"

    # doc.paragraphs[i]._p.addnext(new._element)
exc["序号"]=exc["序号"].fillna("空")
excel=exc[exc["序号"]!="空"]


default_date = '1990/01/01'
excel["验收日期"]=excel["验收日期"].fillna(default_date)
excel["抵达时间"]=excel["抵达时间"].fillna(default_date)
excel["安装时间"]=excel["安装时间"].fillna(default_date)
excel['验收日期'] = pd.to_datetime(excel["验收日期"],errors='coerce').dt.strftime('%Y-%m-%d')
excel['抵达时间'] = pd.to_datetime(excel["抵达时间"],errors='coerce').dt.strftime('%Y-%m-%d')
excel['安装时间'] = pd.to_datetime(excel["安装时间"],errors='coerce').dt.strftime('%Y-%m-%d')
excel["验收日期"] = ['' if i == '1990-01-01' else i
                              for i in excel["验收日期"]]
excel["安装时间"] = ['' if i == '1990-01-01' else i
                              for i in excel["安装时间"]]
excel["抵达时间"] = ['' if i == '1990-01-01' else i
                              for i in excel["抵达时间"]]

excel["客户联系电话："]=excel["客户联系电话："].astype(str)

strl=['序号', '合同编号',  '设备名称', '设备数量', '项目号',  '销售类型',  '客户名称',  '客户联系人']
excel[strl]=excel[strl].fillna("")


doc = docx.Document(doc_name)
# word中第1个表格为doc.tables[0]
table = doc.tables[0]
text = doc.paragraphs[0]


new = deepcopy(table)
doc.add_page_break()
newtext = deepcopy(text)
for i in range(len(excel) - 1):
    new = deepcopy(table)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # paragraph.style = doc.styles['Heading 1']
    text = paragraph.add_run("验收单")

    text.font.size = Pt(24)  # 字体大小
    text.bold = True
    # paragraph.text="验收单"
    paragraph._p.addnext(new._element)
    # doc.paragraphs[i].text="你好"
    doc.add_page_break()
    # doc.paragraphs[i]._p.addnext(new._element)

for i  in range(len(excel)):
    doc.tables[i].cell(0, 2).text =excel["设备名称"][i]
    doc.tables[i].cell(0, 4).text =excel["合同编号"][i]
    doc.tables[i].cell(1, 2).text =excel["项目号"][i]
    doc.tables[i].cell(1, 4).text =excel["设备数量"][i]
    doc.tables[i].cell(2, 4).text =excel["客户名称"][i]
    doc.tables[i].cell(3, 2).text =excel["抵达时间"][i]
    doc.tables[i].cell(3, 4).text =excel["客户联系人"][i]
    doc.tables[i].cell(4, 2).text =excel["安装时间"][i]
    doc.tables[i].cell(4, 4).text =excel["客户联系电话："][i]
    doc.tables[i].cell(17, 4).text =excel["验收日期"][i]
#保存对word表格做的改变
now_time = time.strftime("%Y-%m-%d-%H",time.localtime(time.time()))
book_name='生成验收单'+now_time
doc.save(book_name+'.docx')
end_time = time.time()
print('执行时长:%d秒' % (end_time - start_time))